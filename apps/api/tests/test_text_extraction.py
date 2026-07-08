from pathlib import Path

from docx import Document as DocxDocument

from app.services.text_extraction import extract_text_from_file


PDF_SAMPLE = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 45>>stream
BT /F1 24 Tf 72 720 Td (Hello PDF text) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000234 00000 n 
0000000304 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
399
%%EOF
"""


def test_extracts_txt_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("Hello TXT\nSecond line", encoding="utf-8")

    result = extract_text_from_file(path, source_type="txt", storage_root=tmp_path)

    assert result.success is True
    assert result.source_type == "txt"
    assert "Hello TXT" in result.text
    assert result.metadata["parser"] == "utf-8-text"
    assert result.error is None


def test_extracts_csv_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.csv"
    path.write_text("name,program\nAda,AI\n", encoding="utf-8")

    result = extract_text_from_file(path, source_type="csv", storage_root=tmp_path)

    assert result.success is True
    assert "name, program" in result.text
    assert "Ada, AI" in result.text
    assert result.metadata["row_count"] == 2


def test_extracts_pdf_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    path.write_bytes(PDF_SAMPLE)

    result = extract_text_from_file(path, source_type="pdf", storage_root=tmp_path)

    assert result.success is True
    assert "Hello PDF text" in result.text
    assert result.metadata["parser"] == "pypdf"
    assert result.metadata["page_count"] == 1


def test_extracts_docx_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Hello DOCX text")
    document.add_paragraph("Second paragraph")
    document.save(path)

    result = extract_text_from_file(path, source_type="docx", storage_root=tmp_path)

    assert result.success is True
    assert "Hello DOCX text" in result.text
    assert "Second paragraph" in result.text
    assert result.metadata["parser"] == "python-docx"
    assert result.metadata["paragraph_count"] == 2


def test_rejects_file_outside_storage_root(tmp_path: Path) -> None:
    storage_root = tmp_path / "storage"
    storage_root.mkdir()
    outside_path = tmp_path / "outside.txt"
    outside_path.write_text("outside", encoding="utf-8")

    result = extract_text_from_file(outside_path, source_type="txt", storage_root=storage_root)

    assert result.success is False
    assert result.error is not None
    assert result.error.code == "FILE_OUTSIDE_STORAGE_ROOT"
    assert result.text == ""


def test_unsupported_source_type_returns_structured_error(tmp_path: Path) -> None:
    path = tmp_path / "sample.exe"
    path.write_text("not supported", encoding="utf-8")

    result = extract_text_from_file(path, source_type="exe", storage_root=tmp_path)

    assert result.success is False
    assert result.error is not None
    assert result.error.code == "UNSUPPORTED_SOURCE_TYPE"


def test_parser_failure_returns_structured_error(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a pdf")

    result = extract_text_from_file(path, source_type="pdf", storage_root=tmp_path)

    assert result.success is False
    assert result.error is not None
    assert result.error.code == "EXTRACTION_FAILED"
    assert result.text == ""
